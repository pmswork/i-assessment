"""Generate docs/ELAN_Planner_Manual.docx — the planner user manual.

The generated .docx is committed to the repository so nobody needs to run
this to read the manual. Re-run it after UI changes to keep the manual
current:

    pip install python-docx        # one-off, not a runtime dependency
    python docs/generate_manual.py

The ELAN logo (docs/elan_logo.png) was downloaded once from the official
asset (ImageKit rendering of Logo_Elan Languages_primary.svg as PNG) and is
committed alongside, so regeneration works offline.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

DOCS_DIR = Path(__file__).resolve().parent
LOGO_PATH = DOCS_DIR / "elan_logo.png"
OUTPUT_PATH = DOCS_DIR / "ELAN_Planner_Manual.docx"

ELAN_DARK = RGBColor(0x0C, 0x1F, 0x35)


def _heading(document: Document, text: str, level: int) -> None:
    heading = document.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = ELAN_DARK


def _bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def _steps(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def build_manual() -> None:
    document = Document()

    # -- title page ---------------------------------------------------------
    if LOGO_PATH.exists():
        logo_paragraph = document.add_paragraph()
        logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        logo_paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(3.2))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Interpreter Planner — User Manual")
    title_run.font.size = Pt(26)
    title_run.font.bold = True
    title_run.font.color.rgb = ELAN_DARK

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Automatic interpreter assignment with planner oversight\nMVP edition"
    )
    subtitle_run.font.size = Pt(13)

    document.add_page_break()

    # -- 1. what this application does ---------------------------------------
    _heading(document, "1. What this application does", level=1)
    document.add_paragraph(
        "The ELAN Interpreter Planner assigns interpreters to booked appointments "
        "(court hearings, hospital visits, municipal appointments, remote calls) "
        "automatically, and gives a human planner clear, actionable information for "
        "the small number of jobs the system cannot safely place on its own. The "
        "guiding principle: the system does the routine planning; the planner only "
        "steps in where judgement is genuinely needed."
    )
    _bullets(document, [
        "Automatic assignment respects hard rules that can never be overridden: language match, "
        "sworn qualification where required, availability windows, no double bookings, blacklists, "
        "and physically feasible travel between on-site appointments.",
        "Softer concerns (tight travel, long commutes, cost, workload balance) produce warnings: "
        "the system either handles them itself or asks the planner, depending on the configured autonomy level.",
        "Every job that is not assigned shows exactly why, naming the interpreters that were considered "
        "and the specific obstacle for each — never a vague “no match found”.",
    ])

    # -- 2. getting started ---------------------------------------------------
    _heading(document, "2. Getting started", level=1)
    document.add_paragraph(
        "Ask your administrator for the application address. On a local installation it is "
        "usually http://127.0.0.1:8000. The interface runs in any modern browser; no installation "
        "is needed on the planner's machine."
    )
    document.add_paragraph(
        "The current date and time in Amstelveen is shown as plain text at the top right of "
        "every page, so planners working from elsewhere always reason in Dutch local time."
    )

    # -- 3. the job list -------------------------------------------------------
    _heading(document, "3. The job list (home page)", level=1)
    document.add_paragraph(
        "The home page lists every booking. The three counters at the top show the total, "
        "how many are assigned, and how many still need a human decision. Rows can be filtered "
        "(all / needs decision / assigned) and sorted (date, job, client, language, status, coverage)."
    )
    _bullets(document, [
        "Coverage gauge — a small signal-strength indicator per job: how many qualified interpreters "
        "live within the configured radius. Red or empty means the job sits in a thin part of the map.",
        "Status — the assigned interpreter's name on the left (click it for their profile) and a tag on "
        "the right: “provisional” (proposed by the system, not yet agreed with the interpreter), "
        "“confirmed” (a planner confirmed it, or assigned it by hand), or “auto-confirmed” "
        "(confirmed in bulk on behalf of the interpreters from the Admin page).",
        "The status filter above the list matches these states exactly: All / Needs decision / "
        "Provisional / Auto-confirmed / Confirmed (any).",
        "Urgent — unassigned jobs due within the configured number of days are highlighted in red.",
        "Unassigned jobs show the first reasons directly in the list. Job numbers and interpreter "
        "names inside those reasons are links: click straight through to the job or person mentioned.",
        "Click anywhere on a row to open the job.",
    ])

    # -- 4. provisional vs confirmed -----------------------------------------
    _heading(document, "4. Provisional vs confirmed assignments", level=1)
    document.add_paragraph(
        "Auto-assignment produces proposals. Until someone has actually spoken to the interpreter, "
        "an assignment is only provisional — the system may still replan it when circumstances change "
        "(for example after a re-run of auto-assignment, or when data is edited)."
    )
    _steps(document, [
        "Open the job. A provisional assignment shows an orange “provisional” tag.",
        "Contact the interpreter and agree the job with them.",
        "Click “Confirm with interpreter”. The tag turns to “confirmed”.",
    ])
    document.add_paragraph(
        "Confirmed assignments are planner-owned: re-running auto-assignment never moves them. "
        "If a confirmed assignment later becomes invalid (for example the job time changed and now "
        "overlaps another booking), the system unassigns it and explains why, rather than silently "
        "keeping a broken promise."
    )
    document.add_paragraph(
        "To confirm a whole schedule at once, use “Auto-confirm all provisional” on the Admin page. "
        "Every provisional assignment is promoted to “auto-confirmed” on behalf of the interpreters: "
        "the same planner-owned protection as a normal confirmation, with its own label so it stays "
        "visible that no one individually said yes. Bulk confirmation records no reliability events."
    )

    # -- 5. manual assignment --------------------------------------------------
    _heading(document, "5. Manually assigning a job", level=1)
    _steps(document, [
        "Open the job from the list. The “situation” box summarises what is going on and, "
        "where possible, suggests the best available interpreter including any trade-off.",
        "The candidate table lists only interpreters who meet the basic language and sworn requirements "
        "— anyone else can never take the job, so they are not shown. Blacklisted candidates are marked.",
        "Click “Check” next to a candidate. The system validates the pairing and answers "
        "accepted, warning, or rejected, with the specific reasons.",
        "Accepted: click “Confirm assignment” to save.",
        "Warning: read the reasons and the “Ask before confirming” questions, talk to the "
        "interpreter if needed, then explicitly confirm with “Proceed anyway and confirm”.",
        "Rejected: the assignment cannot be saved. The reasons name the exact obstacle "
        "(for example an overlapping booking, with a link to that job).",
    ])

    # -- 6. recording outcomes ---------------------------------------------------
    _heading(document, "6. Recording outcomes and reliability", level=1)
    document.add_paragraph(
        "After an appointment has taken place, record what happened on the job page: Completed, "
        "No-show, or Late cancellation. These outcomes feed each interpreter's reliability score "
        "(visible on their profile). Reliable interpreters are preferred by auto-assignment over "
        "merely being first or cheapest; no-shows and late cancellations cost points."
    )

    # -- 7. interpreter profiles ---------------------------------------------------
    _heading(document, "7. Interpreter profiles and blacklists", level=1)
    document.add_paragraph(
        "The Interpreters page lists the roster; each profile shows basic details, availability, "
        "the current schedule, the reliability history, and blacklist entries. A blacklist entry "
        "blocks assignment either for one specific client or globally, with an optional reason. "
        "Adding one immediately unassigns any affected jobs — each of those jobs then shows why."
    )

    # -- 8. admin -------------------------------------------------------------------
    _heading(document, "8. Admin: data import, export and editing", level=1)
    _bullets(document, [
        "Import a replacement jobs.csv or interpreters.csv. Files are validated row by row; a single "
        "bad row rejects the whole file with row-numbered reasons, so a half-imported dataset can never occur.",
        "Export the current jobs or roster as CSV at any time; exports can be re-imported unchanged.",
        "Add, edit or delete individual jobs and interpreters with plain forms. Deleting an interpreter "
        "frees their jobs for replanning; editing a job clears its assignment so it is deliberately re-checked.",
        "All admin changes are saved to the local database and survive a restart.",
    ])

    # -- 9. settings -------------------------------------------------------------
    _heading(document, "9. Settings", level=1)
    document.add_paragraph(
        "The Settings page exposes every threshold the system reasons with, including the "
        "auto-assignment autonomy level (how much the system may decide alone), travel speed and "
        "overhead assumptions, warning buffers, the coverage radius, and the urgency window. "
        "Changes apply immediately to the next validation or auto-assignment run and are stored "
        "in the database."
    )

    # -- 10. troubleshooting -----------------------------------------------------
    _heading(document, "10. Troubleshooting", level=1)
    _bullets(document, [
        "A page shows “Not Found” after an update: restart the application server — "
        "a server started before the update keeps serving the old routes.",
        "A job lost its interpreter unexpectedly: open the job; the reasons box explains what "
        "happened (for example a new blacklist entry or an edit that invalidated the pairing).",
        "Numbers look off after changing settings: re-run auto-assignment from the home page to "
        "replan provisional assignments under the new thresholds. Confirmed assignments stay put.",
    ])

    document.save(OUTPUT_PATH)
    print(f"Wrote {OUTPUT_PATH}")


if __name__ == "__main__":
    build_manual()
